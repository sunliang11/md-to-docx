"""Build reference-native.docx for the native engine."""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt

from md_to_docx.paths import data_dir
from md_to_docx.render.fields import add_page_field
from md_to_docx.render.styles import configure_document_styles, set_theme_font_lang


def build_native_reference(output_dir: Path | None = None) -> Path:
    out_dir = output_dir or data_dir()
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "reference-native.docx"

    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(25.4)
    section.right_margin = Mm(25.4)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)

    configure_document_styles(doc, force=True)
    set_theme_font_lang(doc)

    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(fp)

    doc.add_paragraph("Template body placeholder.").style = doc.styles["Normal"]
    doc.save(str(out_path))
    return out_path


def main() -> None:
    path = build_native_reference()
    print(f"Wrote {path}")


if __name__ == "__main__":
    try:
        main()
    except ImportError as exc:
        print("error: python-docx required", file=sys.stderr)
        raise SystemExit(1) from exc
