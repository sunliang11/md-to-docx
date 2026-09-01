"""Build reference-wecom.docx from pandoc default reference template."""

from __future__ import annotations

import shutil
import subprocess
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
except ImportError as exc:
    print(
        "error: python-docx is required to rebuild the reference template.\n"
        "Install it with: pip install python-docx\n"
        "Or install the full dev dependencies: pip install -e \".[dev]\"",
        file=sys.stderr,
    )
    raise SystemExit(1) from exc

from md_to_docx.paths import data_dir

BODY_FONT = "Microsoft YaHei"
MONO_FONT = "Consolas"


def set_fonts(style, latin: str | None = None, east_asia: str | None = None) -> None:
    """Set latin and/or eastAsia fonts for a style."""
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    
    if latin is not None:
        rfonts.set(qn("w:ascii"), latin)
        rfonts.set(qn("w:hAnsi"), latin)
        rfonts.set(qn("w:cs"), latin)
    
    if east_asia is not None:
        rfonts.set(qn("w:eastAsia"), east_asia)


def set_theme_font_lang(doc: Document, latin: str = "en-US", east_asia: str = "zh-CN") -> None:
    """Set theme font language for the document."""
    settings = doc.settings.element
    theme_font_lang = settings.find(qn("w:themeFontLang"))
    if theme_font_lang is None:
        theme_font_lang = OxmlElement("w:themeFontLang")
        settings.append(theme_font_lang)
    theme_font_lang.set(qn("w:val"), latin)
    theme_font_lang.set(qn("w:eastAsia"), east_asia)


def configure_character_style(
    style,
    *,
    latin_font: str | None = None,
    east_asia_font: str | None = None,
    font_size: Pt | None = None,
    bold: bool | None = None,
) -> None:
    if latin_font is not None or east_asia_font is not None:
        set_fonts(style, latin=latin_font, east_asia=east_asia_font)
    if font_size is not None:
        style.font.size = font_size
    if bold is not None:
        style.font.bold = bold


def configure_paragraph_style(
    style,
    *,
    latin_font: str | None = None,
    east_asia_font: str | None = None,
    font_size: Pt | None = None,
    bold: bool | None = None,
    space_before: Pt | None = None,
    space_after: Pt | None = None,
    line_spacing: float | None = None,
    left_indent: Pt | None = None,
) -> None:
    configure_character_style(
        style, latin_font=latin_font, east_asia_font=east_asia_font, font_size=font_size, bold=bold
    )
    pf = style.paragraph_format
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    if line_spacing is not None:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = line_spacing
    if left_indent is not None:
        pf.left_indent = left_indent


def ensure_paragraph_style(doc: Document, name: str, base: str = "Normal") -> None:
    try:
        doc.styles[name]
    except KeyError:
        doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        doc.styles[name].base_style = doc.styles[base]


def ensure_character_style(doc: Document, name: str, base: str = "Default Paragraph Font") -> None:
    try:
        doc.styles[name]
    except KeyError:
        doc.styles.add_style(name, WD_STYLE_TYPE.CHARACTER)
        doc.styles[name].base_style = doc.styles[base]


def ensure_table_style(doc: Document, name: str, base: str = "Table") -> None:
    try:
        doc.styles[name]
    except KeyError:
        doc.styles.add_style(name, WD_STYLE_TYPE.TABLE)
        doc.styles[name].base_style = doc.styles[base]


def add_left_border(style, color: str = "D9D9D9", size: int = 12) -> None:
    ppr = style.element.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), color)
    p_bdr.append(left)
    ppr.append(p_bdr)


def add_table_borders(style, color: str = "000000", size: int = 4) -> None:
    """Add borders to all sides and inside H/V of a table style."""
    tbl_pr = style.element.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        style.element.insert(0, tbl_pr)
    
    tbl_borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tbl_borders.append(border)
    
    existing_borders = tbl_pr.find(qn("w:tblBorders"))
    if existing_borders is not None:
        tbl_pr.remove(existing_borders)
    tbl_pr.append(tbl_borders)


def add_table_header_style(style, bold: bool = True, shading: str = "D9D9D9") -> None:
    """Add header row styling (bold, shading) to a table style."""
    tbl_style_pr = OxmlElement("w:tblStylePr")
    tbl_style_pr.set(qn("w:type"), "firstRow")
    
    rpr = OxmlElement("w:rPr")
    if bold:
        bold_el = OxmlElement("w:b")
        rpr.append(bold_el)
    
    tc_pr = OxmlElement("w:tcPr")
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), shading)
    tc_pr.append(shd)
    
    tbl_style_pr.append(rpr)
    tbl_style_pr.append(tc_pr)
    
    style.element.append(tbl_style_pr)


def configure_table_style(doc: Document) -> None:
    table_style = doc.styles["Table"]
    configure_paragraph_style(
        table_style,
        latin_font=BODY_FONT,
        east_asia_font=BODY_FONT,
        font_size=Pt(10),
        space_before=Pt(0),
        space_after=Pt(0),
    )
    add_table_borders(table_style, color="000000", size=4)
    add_table_header_style(table_style, bold=True, shading="D9D9D9")

    ensure_table_style(doc, "CompactTable", base="Table")
    compact = doc.styles["CompactTable"]
    configure_paragraph_style(
        compact,
        latin_font=BODY_FONT,
        east_asia_font=BODY_FONT,
        font_size=Pt(9),
        space_before=Pt(0),
        space_after=Pt(0),
    )
    add_table_borders(compact, color="000000", size=2)
    add_table_header_style(compact, bold=True, shading="E5E5E5")


def build_reference_doc(output_dir: Path | None = None) -> Path:
    assets = output_dir or data_dir()
    assets.mkdir(parents=True, exist_ok=True)
    default_ref = assets / "reference-default.docx"
    output_ref = assets / "reference-wecom.docx"

    if not default_ref.is_file():
        subprocess.run(
            [
                "pandoc",
                "--print-default-data-file",
                "reference.docx",
                "-o",
                str(default_ref),
            ],
            check=True,
        )

    shutil.copy2(default_ref, output_ref)
    doc = Document(str(output_ref))

    # Set document-wide language
    set_theme_font_lang(doc, latin="en-US", east_asia="zh-CN")

    # Body
    configure_paragraph_style(
        doc.styles["Normal"],
        latin_font=BODY_FONT,
        east_asia_font=BODY_FONT,
        font_size=Pt(11),
        space_after=Pt(6),
        line_spacing=1.5,
    )
    configure_paragraph_style(
        doc.styles["Body Text"],
        latin_font=BODY_FONT,
        east_asia_font=BODY_FONT,
        font_size=Pt(11),
        space_after=Pt(6),
        line_spacing=1.5,
    )
    configure_paragraph_style(
        doc.styles["First Paragraph"],
        latin_font=BODY_FONT,
        east_asia_font=BODY_FONT,
        font_size=Pt(11),
        space_after=Pt(6),
        line_spacing=1.5,
    )
    configure_paragraph_style(
        doc.styles["Compact"],
        latin_font=BODY_FONT,
        east_asia_font=BODY_FONT,
        font_size=Pt(10),
        space_after=Pt(3),
        line_spacing=1.25,
    )

    # Headings
    heading_specs = [
        ("Heading 1", Pt(18), Pt(18), Pt(12)),
        ("Heading 2", Pt(16), Pt(14), Pt(8)),
        ("Heading 3", Pt(14), Pt(12), Pt(6)),
        ("Heading 4", Pt(13), Pt(10), Pt(4)),
        ("Heading 5", Pt(12), Pt(8), Pt(3)),
        ("Heading 6", Pt(11), Pt(6), Pt(3)),
    ]
    for name, size, before, after in heading_specs:
        # Ensure style exists (Heading 5 and 6 may not exist in default reference)
        ensure_paragraph_style(doc, name, base="Normal")
        
        configure_paragraph_style(
            doc.styles[name],
            latin_font=BODY_FONT,
            east_asia_font=BODY_FONT,
            font_size=size,
            bold=True,
            space_before=before,
            space_after=after,
            line_spacing=1.25,
        )
        char_name = f"{name} Char"
        if char_name in [s.name for s in doc.styles]:
            configure_character_style(
                doc.styles[char_name],
                latin_font=BODY_FONT,
                east_asia_font=BODY_FONT,
                font_size=size,
                bold=True
            )

    # Blockquote
    configure_paragraph_style(
        doc.styles["Block Text"],
        latin_font=BODY_FONT,
        east_asia_font=BODY_FONT,
        font_size=Pt(11),
        space_before=Pt(6),
        space_after=Pt(6),
        line_spacing=1.5,
        left_indent=Pt(18),
    )
    add_left_border(doc.styles["Block Text"])

    # Code block
    ensure_paragraph_style(doc, "Source Code", base="Body Text")
    configure_paragraph_style(
        doc.styles["Source Code"],
        latin_font=MONO_FONT,
        east_asia_font=BODY_FONT,
        font_size=Pt(10),
        space_before=Pt(6),
        space_after=Pt(6),
        line_spacing=1.15,
        left_indent=Pt(0),
    )
    doc.styles["Source Code"].font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Inline code
    ensure_character_style(doc, "Verbatim Char")
    set_fonts(doc.styles["Verbatim Char"], latin=MONO_FONT, east_asia=BODY_FONT)
    doc.styles["Verbatim Char"].font.size = Pt(10)
    doc.styles["Verbatim Char"].font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Lists
    if "List Paragraph" in [s.name for s in doc.styles]:
        configure_paragraph_style(
            doc.styles["List Paragraph"],
            latin_font=BODY_FONT,
            east_asia_font=BODY_FONT,
            font_size=Pt(11),
            space_after=Pt(3),
            line_spacing=1.5,
        )

    configure_table_style(doc)

    doc.save(str(output_ref))
    print(f"built: {output_ref}")
    return output_ref


def main() -> int:
    try:
        build_reference_doc()
    except Exception as exc:  # noqa: BLE001
        print(f"error: {exc}", file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
