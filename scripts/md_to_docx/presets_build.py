"""Generate preset template DOCX files."""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

from md_to_docx.paths import assets_dir
from md_to_docx.render.fields import add_page_field
from md_to_docx.render.styles import configure_document_styles, set_theme_font_lang

PRESET_SPECS = {
    "professional": {
        "heading_color": "111827",
        "body_pt": 11,
        "latin": "Calibri",
        "east_asia": "Microsoft YaHei",
        "header": None,
    },
    "technical": {
        "heading_color": "1E3A5F",
        "body_pt": 10.5,
        "latin": "Calibri",
        "east_asia": "Microsoft YaHei",
        "header": "Document Title",
    },
    "academic": {
        "heading_color": "000000",
        "body_pt": 12,
        "latin": "Times New Roman",
        "east_asia": "SimSun",
        "header": "Paper Title",
    },
    "business": {
        "heading_color": "1F4E79",
        "body_pt": 11,
        "latin": "Calibri",
        "east_asia": "Microsoft YaHei",
        "header": "Company Report",
    },
    "report": {
        "heading_color": "0F172A",
        "body_pt": 11,
        "latin": "Calibri",
        "east_asia": "Microsoft YaHei",
        "header": "Report",
    },
}


def _hex_color(hex_str: str) -> RGBColor:
    return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))


def build_preset(name: str, spec: dict, out_dir: Path) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(25.4)
    section.right_margin = Mm(25.4)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)

    configure_document_styles(doc, force=True)
    normal = doc.styles["Normal"]
    from md_to_docx.render.styles import _set_fonts

    _set_fonts(normal, latin=spec["latin"], east_asia=spec["east_asia"])
    normal.font.size = Pt(spec["body_pt"])
    for level in range(1, 7):
        h = doc.styles[f"Heading {level}"]
        _set_fonts(h, latin=spec["latin"], east_asia=spec["east_asia"])
        h.font.color.rgb = _hex_color(spec["heading_color"])

    set_theme_font_lang(doc)
    if spec.get("header"):
        hp = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
        hp.text = spec["header"]
    fp = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(fp)

    doc.add_paragraph(f"{name} template")
    out = out_dir / f"{name}.docx"
    doc.save(str(out))
    return out


def main() -> None:
    out_dir = assets_dir() / "presets"
    out_dir.mkdir(parents=True, exist_ok=True)
    for name, spec in PRESET_SPECS.items():
        path = build_preset(name, spec, out_dir)
        print(f"Wrote {path}")
    native = assets_dir() / "reference-native.docx"
    from md_to_docx.reference_native import build_native_reference

    build_native_reference(assets_dir())
    print(f"Wrote {native}")


if __name__ == "__main__":
    main()
