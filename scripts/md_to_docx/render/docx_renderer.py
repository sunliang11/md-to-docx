"""AST to DOCX renderer."""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

from md_to_docx.ast import nodes as n
from md_to_docx.render.fields import add_toc_field
from md_to_docx.render.header_footer import apply_header_footer
from md_to_docx.render.styles import configure_document_styles
from md_to_docx.render.template import open_document

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


def _add_hyperlink(paragraph: Paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    new_run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _apply_list_level(paragraph: Paragraph, ordered: bool, level: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(min(level, 8)))
    numpr.append(ilvl)
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), "1" if ordered else "2")
    numpr.append(numid)
    ppr.append(numpr)


class DocxRenderer:
    def __init__(
        self,
        doc: Document,
        *,
        base_dir: Path | None = None,
        figure_label: str = "Figure",
        table_label: str = "Table",
        section_label: str = "Section",
        xref_map: dict[str, tuple[str, int]] | None = None,
        media_paths: dict[str, Path] | None = None,
    ) -> None:
        self.doc = doc
        self.base_dir = base_dir or Path.cwd()
        self.figure_label = figure_label
        self.table_label = table_label
        self.section_label = section_label
        self.xref_map = xref_map or {}
        self.media_paths = media_paths or {}
        self._bookmark_id = 0

    def _next_bookmark_id(self) -> int:
        self._bookmark_id += 1
        return self._bookmark_id

    def render(self, document: n.Document) -> None:
        meta = document.metadata
        if meta.title and document.blocks and not (
            isinstance(document.blocks[0], n.Heading) and document.blocks[0].level == 1
        ):
            self.doc.add_heading(meta.title, level=1)

        for block in document.blocks:
            self.render_block(block)

    def render_block(self, block: n.Block) -> None:
        if isinstance(block, n.Heading):
            self._render_heading(block)
        elif isinstance(block, n.Paragraph):
            self._render_paragraph(block)
        elif isinstance(block, n.ListBlock):
            self._render_list(block, level=0)
        elif isinstance(block, n.Table):
            self._render_table(block)
        elif isinstance(block, n.CodeBlock):
            self._render_code(block)
        elif isinstance(block, n.Mermaid):
            self._render_mermaid(block)
        elif isinstance(block, n.MathBlock):
            self._render_math_block(block)
        elif isinstance(block, n.BlockQuote):
            self._render_blockquote(block)
        elif isinstance(block, n.ThematicBreak):
            p = self.doc.add_paragraph()
            p.paragraph_format.border_bottom = True
        elif isinstance(block, n.Image):
            self._render_image(block)
        elif isinstance(block, n.Figure):
            self._render_figure(block)
        elif isinstance(block, n.PageBreak):
            self.doc.add_page_break()
        elif isinstance(block, n.TableOfContents):
            self._render_toc(block)
        elif isinstance(block, n.HTMLBlock):
            self.doc.add_paragraph(block.raw)

    def _render_heading(self, block: n.Heading) -> None:
        p = self.doc.add_heading(level=block.level)
        self._render_inlines(p, block.children)

    def _render_paragraph(self, block: n.Paragraph) -> None:
        p = self.doc.add_paragraph()
        self._render_inlines(p, block.children)

    def _render_inlines(self, paragraph: Paragraph, children: tuple[n.Inline, ...]) -> None:
        for child in children:
            if isinstance(child, n.Text):
                paragraph.add_run(child.value)
            elif isinstance(child, n.Strong):
                run = paragraph.add_run()
                run.bold = True
                for sub in child.children:
                    if isinstance(sub, n.Text):
                        run.text += sub.value
            elif isinstance(child, n.Emphasis):
                run = paragraph.add_run()
                run.italic = True
                for sub in child.children:
                    if isinstance(sub, n.Text):
                        run.text += sub.value
            elif isinstance(child, n.Strike):
                run = paragraph.add_run()
                run.font.strike = True
                for sub in child.children:
                    if isinstance(sub, n.Text):
                        run.text += sub.value
            elif isinstance(child, n.Code):
                run = paragraph.add_run(child.value)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            elif isinstance(child, n.Link):
                text = "".join(
                    c.value for c in child.children if isinstance(c, n.Text)
                ) or child.href
                _add_hyperlink(paragraph, text, child.href)
            elif isinstance(child, n.InlineImage):
                self._add_inline_image(paragraph, child)
            elif isinstance(child, n.Break):
                paragraph.add_run().add_break(WD_BREAK.LINE)
            elif isinstance(child, n.SoftBreak):
                paragraph.add_run().add_break(WD_BREAK.LINE)
            elif isinstance(child, n.MathInline):
                from md_to_docx.render.omml import add_inline_math

                add_inline_math(paragraph, child.latex)
            elif isinstance(child, n.CrossRef):
                label, num = self.xref_map.get(
                    f"{child.kind}:{child.identifier}", ("?", 0)
                )
                text = f"{label} {num}"
                _add_hyperlink(paragraph, text, f"#{child.kind}-{child.identifier}")
            elif isinstance(child, n.FootnoteRef):
                from md_to_docx.render.footnotes import add_footnote_ref

                add_footnote_ref(paragraph, child.key)

    def _render_list(self, block: n.ListBlock, level: int) -> None:
        style = "List Number" if block.ordered else "List Bullet"
        for item in block.items:
            for sub in item.children:
                if isinstance(sub, n.ListBlock):
                    self._render_list(sub, level + 1)
                elif isinstance(sub, n.Paragraph):
                    p = self.doc.add_paragraph(style=style)
                    _apply_list_level(p, block.ordered, level)
                    self._render_inlines(p, sub.children)
                else:
                    self.render_block(sub)

    def _render_table(self, block: n.Table) -> None:
        if not block.rows:
            return
        rows = len(block.rows)
        cols = max(len(r) for r in block.rows)
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"
        for ri, row in enumerate(block.rows):
            for ci, cell in enumerate(row):
                tc = table.cell(ri, ci)
                tc.text = ""
                for sub in cell.children:
                    if isinstance(sub, n.Paragraph):
                        p = tc.paragraphs[0] if tc.paragraphs else tc.add_paragraph()
                        self._render_inlines(p, sub.children)
                    else:
                        self.render_block(sub)
                if cell.header or ri == 0:
                    for p in tc.paragraphs:
                        for run in p.runs:
                            run.bold = True

    def _render_code(self, block: n.CodeBlock) -> None:
        for line in block.text.splitlines() or [""]:
            p = self.doc.add_paragraph(line, style="MDCodeBlock")

    def _render_mermaid(self, block: n.Mermaid) -> None:
        from md_to_docx.render.image import embed_image_path

        key = f"mermaid:{hash(block.source)}"
        path = self.media_paths.get(key)
        if path and path.is_file():
            embed_image_path(self.doc, path)
        else:
            self._render_code(n.CodeBlock(block.source, lang="mermaid"))

    def _render_math_block(self, block: n.MathBlock) -> None:
        from md_to_docx.render.omml import add_block_math

        p = self.doc.add_paragraph()
        add_block_math(p, block.latex)

    def _render_blockquote(self, block: n.BlockQuote) -> None:
        for sub in block.children:
            if isinstance(sub, n.Paragraph):
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
                self._render_inlines(p, sub.children)
            else:
                self.render_block(sub)

    def _resolve_image(self, src: str) -> Path | None:
        path = Path(src)
        if not path.is_absolute():
            path = self.base_dir / path
        return path if path.is_file() else None

    def _render_image(self, block: n.Image) -> None:
        from md_to_docx.render.image import embed_image_path

        resolved = self._resolve_image(block.src)
        if resolved:
            embed_image_path(self.doc, resolved, alt=block.alt)
        else:
            print(f"warning: missing image: {block.src}", file=sys.stderr)
            self.doc.add_paragraph(f"[missing image: {block.src}]")

    def _render_figure(self, block: n.Figure) -> None:
        bid = self._next_bookmark_id()
        if isinstance(block.image, n.Image):
            self._render_image(block.image)
        elif isinstance(block.image, n.Mermaid):
            self._render_mermaid(block.image)
        num = block.number or 0
        cap = self.doc.add_paragraph(style="Caption")
        cap.add_run(f"{self.figure_label} {num}. {block.caption}")
        p = cap._p
        bm_start = OxmlElement("w:bookmarkStart")
        bm_start.set(qn("w:id"), str(bid))
        bm_start.set(qn("w:name"), f"fig-{block.identifier}")
        p.insert(0, bm_start)
        bm_end = OxmlElement("w:bookmarkEnd")
        bm_end.set(qn("w:id"), str(bid))
        p.append(bm_end)

    def _render_toc(self, block: n.TableOfContents) -> None:
        self.doc.add_heading(block.title, level=1)
        p = self.doc.add_paragraph()
        add_toc_field(p)

    def _add_inline_image(self, paragraph: Paragraph, image: n.InlineImage) -> None:
        resolved = self._resolve_image(image.src)
        if resolved:
            run = paragraph.add_run()
            run.add_picture(str(resolved), width=Inches(4))


def render_docx(
    document: n.Document,
    out_path: Path,
    *,
    base_dir: Path | None = None,
    template_path: Path | None = None,
    title: str | None = None,
    author: str | None = None,
    date: str | None = None,
    version: str | None = None,
    page_numbers: bool = True,
    figure_label: str = "Figure",
    table_label: str = "Table",
    section_label: str = "Section",
    xref_map: dict[str, tuple[str, int]] | None = None,
    media_paths: dict[str, Path] | None = None,
) -> None:
    doc = open_document(template_path)
    configure_document_styles(doc)
    meta = document.metadata
    apply_header_footer(
        doc,
        title=title or meta.title,
        author=author or meta.author,
        date=date or meta.date,
        version=version or meta.version,
        page_numbers=page_numbers,
    )
    renderer = DocxRenderer(
        doc,
        base_dir=base_dir,
        figure_label=figure_label,
        table_label=table_label,
        section_label=section_label,
        xref_map=xref_map,
        media_paths=media_paths,
    )
    renderer.render(document)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
