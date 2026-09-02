"""Image embedding helpers."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches


def embed_image_path(doc: Document, path: Path, *, alt: str = "", width_inches: float = 5.0) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    if alt:
        p.alignment = 0
