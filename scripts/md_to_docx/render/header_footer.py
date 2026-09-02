"""Header and footer rendering."""

from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from md_to_docx.render.fields import add_page_field


def apply_header_footer(
    doc: Document,
    *,
    title: str | None = None,
    author: str | None = None,
    date: str | None = None,
    version: str | None = None,
    page_numbers: bool = True,
) -> None:
    for section in doc.sections:
        header = section.header
        footer = section.footer

        if title or author:
            header.paragraphs[0].clear() if header.paragraphs else header.add_paragraph()
            hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            if title:
                hp.add_run(title)
            if author:
                if title:
                    hp.add_run("\t")
                run = hp.add_run(author)
                run.bold = False
                hp.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if page_numbers or version or date:
            fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            if version:
                fp.add_run(version)
            if date:
                if version:
                    fp.add_run("  ")
                fp.add_run(date)
            if page_numbers:
                fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if version or date:
                    fp.add_run("\t")
                add_page_field(fp)
