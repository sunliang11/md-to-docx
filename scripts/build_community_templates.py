#!/usr/bin/env python3
"""Build community template assets (template.docx + preview.png)."""

from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
TEMPLATES = ROOT / "templates"
sys.path.insert(0, str(ROOT / "scripts"))

from md_to_docx.presets_build import PRESET_SPECS, _hex_color, build_preset  # noqa: E402
from md_to_docx.render.fields import add_page_field  # noqa: E402
from md_to_docx.render.styles import _set_fonts, configure_document_styles, set_theme_font_lang  # noqa: E402

TEMPLATE_SPECS = {
    "technical-design": {
        "preset": "technical",
        "header": "Technical Design",
        "preview_color": (30, 58, 95),
        "preview_title": "Technical Design",
    },
    "consulting-report": {
        "preset": "business",
        "header": "Consulting Report",
        "preview_color": (31, 78, 121),
        "preview_title": "Consulting Report",
    },
    "academic-ieee-ish": {
        "preset": "academic",
        "header": "Conference Paper",
        "preview_color": (0, 0, 0),
        "preview_title": "Academic Paper",
    },
    "chinese-official": {
        "preset": "report",
        "header": "公文标题",
        "preview_color": (15, 23, 42),
        "preview_title": "公文模板",
        "east_asia": "SimSun",
    },
}


def _build_custom_template(name: str, spec: dict, out_dir: Path) -> Path:
  preset_name = spec["preset"]
  preset_spec = dict(PRESET_SPECS[preset_name])
  if "east_asia" in spec:
    preset_spec["east_asia"] = spec["east_asia"]
  preset_spec["header"] = spec["header"]

  doc = Document()
  section = doc.sections[0]
  section.page_width = Mm(210)
  section.page_height = Mm(297)
  section.left_margin = Mm(25.4)
  section.right_margin = Mm(25.4)
  section.top_margin = Mm(25.4)
  section.bottom_margin = Mm(25.4)

  configure_document_styles(doc, force=True)
  normal = doc.styles["Normal"]
  _set_fonts(normal, latin=preset_spec["latin"], east_asia=preset_spec["east_asia"])
  normal.font.size = Pt(preset_spec["body_pt"])
  for level in range(1, 7):
    h = doc.styles[f"Heading {level}"]
    _set_fonts(h, latin=preset_spec["latin"], east_asia=preset_spec["east_asia"])
    h.font.color.rgb = _hex_color(preset_spec["heading_color"])

  set_theme_font_lang(doc)
  hp = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
  hp.text = spec["header"]
  fp = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
  fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
  add_page_field(fp)

  doc.add_paragraph(f"{name} template")
  out = out_dir / "template.docx"
  doc.save(str(out))
  return out


def _build_preview(name: str, spec: dict, out_dir: Path) -> Path:
  from PIL import Image, ImageDraw, ImageFont

  width, height = 400, 520
  color = spec["preview_color"]
  img = Image.new("RGB", (width, height), (255, 255, 255))
  draw = ImageDraw.Draw(img)
  draw.rectangle([0, 0, width, 80], fill=color)
  try:
    font = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 22)
    subfont = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 14)
  except OSError:
    font = ImageFont.load_default()
    subfont = font
  draw.text((20, 28), spec["preview_title"], fill=(255, 255, 255), font=font)
  draw.text((20, 120), f"Template: {name}", fill=(60, 60, 60), font=subfont)
  draw.text((20, 150), "md-to-docx community template", fill=(120, 120, 120), font=subfont)
  draw.rectangle([20, 200, width - 20, 480], outline=(200, 200, 200), width=1)
  out = out_dir / "preview.png"
  img.save(out)
  return out


def main() -> None:
  license_src = TEMPLATES / "MIT-LICENSE.txt"
  for name, spec in TEMPLATE_SPECS.items():
    out_dir = TEMPLATES / name
    out_dir.mkdir(parents=True, exist_ok=True)
    path = _build_custom_template(name, spec, out_dir)
    print(f"Wrote {path}")
    preview = _build_preview(name, spec, out_dir)
    print(f"Wrote {preview}")
    if license_src.is_file():
      shutil.copy(license_src, out_dir / "LICENSE")


if __name__ == "__main__":
  main()
